import datetime
import os
from tkinter import *
from docx import Document
    
last_date = None

# Create the main window
root = Tk()

# Set the window title
root.title("JustJot")

# Set the window size
root.geometry("500x700+{}+{}".format(int(root.winfo_screenwidth()*0.7), int(root.winfo_screenheight()*0.1)))

# Create a label for the title
title_label = Label(root, text="JustJot", font=("Italic", 20))
title_label.pack(pady=20)

# Create a label for the journal entry
entry_label = Label(root, text="Title (optional):", font=("Helvetica", 12), anchor='w')
entry_label.pack(padx=20, pady=10)

# Create a text box for the title
title_text = Entry(root, width=60, font=("Helvetica", 14))
title_text.pack(padx=20, pady=10)

# Create a text box for the journal entry
entry_text = Text(root, width=60, height=25, font=("Helvetica", 14), wrap=WORD)
entry_text.pack(padx=20, pady=10)

# Add a scrollbar to the text box
scrollbar = Scrollbar(root, command=entry_text.yview)
scrollbar.pack(side=RIGHT, fill=Y)
entry_text.config(yscrollcommand=scrollbar.set)

# Set the scrollbar to fill the y-axis of the window
scrollbar.pack(side=RIGHT, fill=Y)

#documents
home_dir = os.path.expanduser("~")
file_name = str(datetime.date.today().year) + 'Jots.docx'
file_path = 'E:\\' + file_name

if os.path.isfile(file_path):
    journal = Document(file_path)
else:
    journal = Document()

def on_key_press(event):
    global last_date

    if event.keysym == 'Return':
        if event.state == 12:
            value=entry_text.get("1.0",END)
            if not value:
                return

            if title_text.get() != '':
                isTitled = True
            else:
                isTitled = False

            now = datetime.datetime.now()
            dt_string = now.strftime("%d/%m/%Y")

                
            time = now.strftime('%H:%M')
            dated = False
            for para in reversed(journal.paragraphs):
                if dt_string in para.text:
                    dated = True

            if not dated:
                journal.add_paragraph('====================================================================\n' + dt_string)
            if isTitled:
                p = journal.add_paragraph(time)
                p.add_run(f'\n{title_text.get()}').bold = True
                p.add_run(f'\n\t{value}')
            else:
                journal.add_paragraph(f'{time}\n{value}')
            title_text.delete(0, END)
            entry_text.delete('1.0', END)
            print(value)
            journal.save(file_path)

root.bind("<Control-KeyPress>", on_key_press)
title_text.focus_set()



# Run the main event loop
root.mainloop()

